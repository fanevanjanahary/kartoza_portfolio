import frappe
from frappe.utils.pdf import get_pdf
from frappe import _
import io
import requests
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from bs4 import BeautifulSoup
from docx.shared import Inches, Pt
from datetime import datetime
import tempfile
import os
import re
from html import unescape
import zipfile
from io import BytesIO

@frappe.whitelist()
def export_portfolio(portfolio_names, format, layout):
    if not portfolio_names:
        frappe.throw(_("No portfolio names provided"))

    file_data_list = []
    if layout == 'kartoza':
        content = generate_kartoza_html_content(portfolio_names)
    elif layout == 'world bank':
        content = worldbank_format_html(portfolio_names)

    if format == "pdf":
        file_data = get_pdf(content)
        file_extension = "pdf"
    elif format == "docx":
        file_data = generate_docx_from_html(content)
        file_extension = "docx"
    elif format == "html":
        html_file_data = generate_html_file(content)
        file_data_list.append((html_file_data, "html"))
        file_extension = "zip"
        # Create a ZIP file
        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, 'w') as zip_file:
            timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
            for file_data, ext in file_data_list:
                file_name = f"portfolio_export_{timestamp}.{ext}"
                zip_file.writestr(file_name, file_data)
        zip_buffer.seek(0)
        file_data = zip_buffer.getvalue()
    else:
        frappe.throw(_("Unsupported file format"))

    if format != "html":
        file_data_list.append((file_data, file_extension))

        # Generate a default file name with timestamp
        timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
        default_file_name = f"portfolio_export_{timestamp}.{file_extension}"

        # Create a new File document to save the generated file
        file_doc = frappe.get_doc({
            "doctype": "File",
            "file_name": default_file_name,
            "is_private": 1,
            "content": file_data
        })
        file_doc.insert()

        return {
            "status": "success",
            "message": f"Portfolios exported successfully.",
            "file_url": file_doc.file_url
        }
    else:
        # For HTML format, handle the ZIP file case
        zip_file_name = f"portfolio_export_{timestamp}.zip"
        file_doc = frappe.get_doc({
            "doctype": "File",
            "file_name": zip_file_name,
            "is_private": 1,
            "content": file_data
        })
        file_doc.insert()

        return {
            "status": "success",
            "message": f"Portfolios exported successfully.",
            "file_url": file_doc.file_url
        }


def generate_kartoza_html_content(portfolios):
    project_details = ""
    portfolio_names = frappe.parse_json(portfolios)
    for docname in portfolio_names:
        portfolio = frappe.get_doc("Portfolio", docname)
        technologies_list = ""
        for tech in portfolio.technologies:
            technologies_list += f"<li>{tech.technology}</li>"
        
        client_contact = portfolio.contact if portfolio.contact != "" else "Unavailable"
        client_reference = portfolio.client_reference if portfolio.client_reference != "" else "Unavailable"
        client_logo = frappe.utils.get_url() + portfolio.client_logo if portfolio.client_logo  else ""

        services_list = ""
        for service in portfolio.services_listed:
            services_list += f"<li>{service.service}</li>"

        absolute_url = frappe.utils.get_url() 
        time = absolute_url + "/assets/portfolio/images/time.png"
        location = absolute_url + "/assets/portfolio/images/location.png"
        person = absolute_url + "/assets/portfolio/images/person.png"
        footer = absolute_url + "/assets/portfolio/images/footer.png"

        images_list = ""
        for image in portfolio.images:
            if image and image.website_image:
                image_url = image.website_image
                if not image_url.startswith(('http://', 'https://')):
                    image_url = frappe.utils.get_url() + image_url
                images_list += f'<img src="{image_url}" alt="Screenshot" style="width:100%; height:auto; object-fit:contain; padding:10px;"><br>'

        project_details += f"""
        <div style="page-break-after:always">
            <div style="height:100%; position:relative">
                <h3 style="color:#f4b340; text-align:center;">Kartoza Project Sheet</h3>
                <h2 style="text-align:center;">{portfolio.title}</h2>
                <div>
                    <hr style="border: 8px solid #f4b340; width: 90px; margin:auto;">
                </div>
                <br><br>
                <table style="width:100%; border-collapse:collapse;">
                    <tr>
                        <td style="width:33%; text-align:center; border:1px solid gray; padding:10px;">
                            <div>
                                <img src="{person}" alt="Project Image" style="width:80px; height:auto;">
                                <p>Client: {portfolio.client}</p>
                            </div>
                        </td>
                        <td style="width:33%; text-align:center; border:1px solid gray; padding:10px;">
                            <img src="{location}" alt="Project Image" style="width:80px; height:auto;">
                            <p>Location: {portfolio.location}</p>
                        </td>
                        <td style="width:33%; text-align:center; border:1px solid gray; padding:10px;">
                            <img src="{time}" alt="Project Image" style="width:80px; height:auto;">
                            <p>Period: {portfolio.start_date} - {portfolio.end_date}</p>
                        </td>
                    </tr>
                </table>
                <table style="width:100%; border-collapse:collapse;">
                    <tr>
                        <td style="width:40%; border:1px solid gray; vertical-align:top; padding:10px;">
                            <div style="width:100%; height:100px; border:1px solid gray;">
                                <img src="{client_logo}" style="width:100%;height:100px; object-fit:contain;"/>
                            </div>
                            <div style="width:100%; height:100px; border:1px solid gray;">
                                Client reference: {client_reference}
                            </div>
                            <div style="width:100%; height:100px; border:1px solid gray;">
                                Client contact: {client_contact}
                            </div>
                        </td>
                        <td style="width:60%; border:1px solid gray; vertical-align:top; padding:10px;">
                            <div style="width:100%;overflow:hidden;">
                                {images_list}
                            </div>
                        </td>
                    </tr>
                </table>
                <table style="width:100%; border-collapse:collapse;">
                    <tr>
                        <td style="width:55%; border:1px solid gray; padding:10px;">
                            <p>Project Description</p>
                            <p>{portfolio.body}</p>
                        </td>
                        <td style="width:45%; border:1px solid gray; padding:10px;vertical-align: top;">
                            <div>
                                <p>Services Provided</p>
                                <ul>
                                    {services_list}
                                </ul>
                            </div>
                        </td>
                    </tr>
                </table>
                <div>
                    <img src="{footer}" alt="Project Image" style="width:100%; height:220px; text-align:center; position:absolute; bottom:8px; left:0;">
                </div>
            </div>
        </div>
        """

    content = f"""
    <html>
    <head>
        <title>Kartoza Project Sheet</title>
    </head>
    <body>
        {project_details}
        
    </body>
    </html>
    """
    return content


def generate_html_file(content):
    absolute_url = frappe.utils.get_url() 
    footer = absolute_url + "/assets/portfolio/images/footer.png"
    str_to_remove = f'<img src="{footer}" alt="Project Image" style="width:100%; height:220px; text-align:center; position:absolute; bottom:8px; left:0;">'
    result = content.replace(str_to_remove, "")
    result = result.replace('<div style="page-break-after:always">', '<div style="page-break-after:always;height:110%">')
    output = io.BytesIO()
    output.write(result.encode('utf-8'))
    output.seek(0)
    return output.getvalue()


def strip_html_tags(text):
    """Remove HTML tags from a string."""
    clean = re.compile('<.*?>')
    return re.sub(clean, '', unescape(text))



def generate_docx_from_html(html_content):
    # Parse the HTML content
    soup = BeautifulSoup(html_content, 'html.parser')

    # Create a new Document
    doc = Document()

    # Add the title (h3)
    title = soup.find('h3')
    if title:
        doc.add_heading(title.text.strip(), level=1)

    # Add the subtitle (h2)
    subtitle = soup.find('h2')
    if subtitle:
        doc.add_heading(subtitle.text.strip(), level=2)

    # Find all tables and add them to the document
    tables = soup.find_all('table')
    for table in tables:
        rows = table.find_all('tr')
        if rows:
            first_row_cells = rows[0].find_all('td')
            doc_table = doc.add_table(rows=len(rows), cols=len(first_row_cells))
            for i, row in enumerate(rows):
                cells = row.find_all('td')
                for j, cell in enumerate(cells):
                    cell_content = cell.get_text(strip=True)
                    
                    # Only add text content from the cell if there's no <ul> inside it
                    ul = cell.find('ul')
                    if not ul:
                        doc_table.cell(i, j).text = cell_content

                    # Handle images inside table cells
                    img_tag = cell.find('img')
                    if img_tag and img_tag.get('src'):
                        image_url = img_tag['src']
                        try:
                            response = requests.get(image_url)
                            image_data = BytesIO(response.content)
                            doc_table.cell(i, j).add_paragraph().add_run().add_picture(image_data, width=Inches(1.5))
                        except requests.RequestException:
                            print(f"Failed to download image from {image_url}")

                    # Check for <ul> inside the cell and handle it separately
                    if ul:
                        for li in ul.find_all('li'):
                            # Add each list item as a bullet point in the same cell
                            doc_table.cell(i, j).add_paragraph(li.get_text(strip=True), style='ListBullet')

    # Add the footer image
    # footer_img_tag = soup.find('img', alt='Project Image')
    # if footer_img_tag and footer_img_tag.get('src'):
    #     footer_img_url = footer_img_tag['src']
    #     try:
    #         response = requests.get(footer_img_url)
    #         footer_image_data = BytesIO(response.content)
    #         doc.add_picture(footer_image_data, width=Inches(6))  # Adjust size as needed
    #     except requests.RequestException:
    #         print(f"Failed to download footer image from {footer_img_url}")

    # Save the document
    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()

def generate_docx_content(portfolios):
    document = Document()
    
    # Define image paths
    base_url = frappe.utils.get_url()
    image_paths = {
        'time': f"{base_url}/assets/portfolio/images/time.png",
        'location': f"{base_url}/assets/portfolio/images/location.png",
        'person': f"{base_url}/assets/portfolio/images/person.png",
        'footer': f"{base_url}/assets/portfolio/images/footer.png"
    }
    
    # Parse JSON
    portfolio_names = frappe.parse_json(portfolios)
    
    for docname in portfolio_names:
        portfolio = frappe.get_doc("Portfolio", docname)

        # Create a title
        document.add_heading('Kartoza Project Sheet', level=2).alignment = 1  # Center alignment
        document.add_heading(portfolio.title, level=1).alignment = 1  # Center alignment

        # Add a horizontal line
        document.add_paragraph().add_run().add_break()
        p = document.add_paragraph()
        p.add_run().add_break()
        p.add_run().add_break()
        
        # Add project information table
        table = document.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Client'
        hdr_cells[1].text = 'Location'
        hdr_cells[2].text = 'Period'
        
        row_cells = table.add_row().cells
        row_cells[0].text = portfolio.client
        row_cells[1].text = portfolio.location
        row_cells[2].text = f"{portfolio.start_date} - {portfolio.end_date}"

        # Add client details table
        document.add_paragraph().add_run().add_break()
        table = document.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Client Reference'
        hdr_cells[1].text = 'Client Contact'
        
        row_cells = table.add_row().cells
        row_cells[0].text = portfolio.client_reference if portfolio.client_reference else 'Unavailable'
        row_cells[1].text = portfolio.contact if portfolio.contact else 'Unavailable'

        # Add project description and services
        document.add_heading('Project Description', level=2)
        description_text = strip_html_tags(portfolio.body)
        document.add_paragraph(description_text)
        
        document.add_heading('Services Provided', level=2)
        services_list = '\n'.join([service.service for service in portfolio.services_listed])
        document.add_paragraph(services_list)

        # Add images
        if portfolio.images:
            document.add_heading('Project Images', level=2)
            for image in portfolio.images:
                if image and image.website_image:
                    image_url = image.website_image
                    # Check if the URL is missing a schema and prepend one if necessary
                    if not image_url.startswith(('http://', 'https://')):
                        image_url = frappe.utils.get_url() + image_url
                    try:
                        response = requests.get(image_url)
                        response.raise_for_status()  # Check if the request was successful
                        image_stream = io.BytesIO(response.content)
                        document.add_picture(image_stream, width=Inches(5))  # Adjust width as needed
                        document.add_paragraph().add_run().add_break()
                    except requests.RequestException:
                        print(f"Could not fetch image from {image_url}")

        # Add footer image
        footer_image_url = image_paths['footer']
        try:
            response = requests.get(footer_image_url)
            response.raise_for_status()  # Check if the request was successful
            footer_image_stream = io.BytesIO(response.content)
            document.add_picture(footer_image_stream, width=Inches(6))
        except requests.RequestException:
            print(f"Could not fetch footer image from {footer_image_url}")

        # Add page break after each portfolio
        document.add_page_break()
    
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def worldbank_format_html(portfolios):
    """Create a World Bank format HTML document for the given portfolios."""
    portfolio_names = frappe.parse_json(portfolios)
    
    # Initialize the HTML string
    html = """
    <!DOCTYPE html>
    <html lang="en">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>World Bank Format Document</title>
        <style>
            body {
                font-family: Arial, sans-serif;
            }
            h1, h2 {
                font-weight: bold;
            }
            table {
                border-collapse: collapse;
                width: 100%;
            }
            table, th, td {
                border: 1px solid black;
            }
            th, td {
                padding: 8px;
                text-align: left;
            }
            th {
                width: 200px;
            }
            td {
                width: 300px;
            }
        </style>
    </head>
    <body>
        <h1>Assignment Details</h1>
    """

    # Loop through each portfolio and generate the HTML content
    for portfolio_name in portfolio_names:
        details = frappe.get_doc("Portfolio", portfolio_name)
        
        html += f"""
        <h2>{details.title}</h2>
        <table>
            <tr>
                <th>Assignment name:</th>
                <td>{details.title}</td>
            </tr>
            <tr>
                <th>Approx. value of the contract (in current US$):</th>
                <td>{details.approximate_contract_value}</td>
            </tr>
            <tr>
                <th>Country:</th>
                <td>{details.location}</td>
            </tr>
            <tr>
                <th>Duration of assignment (months):</th>
                <td>{details.duration_of_assignment}</td>
            </tr>
            <tr>
                <th>Name of Client(s):</th>
                <td>{details.client}</td>
            </tr>
            <tr>
                <th>Contact Person, Title/Designation, Tel. No./Address:</th>
                <td>{details.contact}</td>
            </tr>
            <tr>
                <th>Start Date (month/year):</th>
                <td>{details.start_date}</td>
            </tr>
            <tr>
                <th>End Date (month/year):</th>
                <td>{details.end_date}</td>
            </tr>
            <tr>
                <th>Total No. of staff-months of the assignment:</th>
                <td>{details.total_staff_months}</td>
            </tr>
            <tr>
                <th>No. of professional staff-months provided by your consulting firm/organization or your sub consultants:</th>
                <td>{details.total_staff_months}</td>
            </tr>
            <tr>
                <th>Name of associated Consultants, if any:</th>
                <td></td>
            </tr>
            <tr>
                <th>Name of senior professional staff of your consulting firm/organization involved and designation and/or functions performed:</th>
                <td></td>
            </tr>
            <tr>
                <th>Description of Project:</th>
                <td>{details.body}</td>
            </tr>
            <tr>
                <th>Description of actual services provided by your staff within the assignment:</th>
                <td>{details.services_listed}</td>
            </tr>
        </table>
        """

    # Close the HTML tags
    html += """
    </body>
    </html>
    """

    return html


def worldbank_format(portfolios):
    """Create a World Bank format document for the given portfolios."""
    portfolio_names = frappe.parse_json(portfolios)
    doc = Document()

    # Add Title
    title = doc.add_heading(level=1)
    title_run = title.add_run("Assignment Details")
    title_run.bold = True

    # Loop through each portfolio and create a table
    for portfolio_name in portfolio_names:
        details = frappe.get_doc("Portfolio", portfolio_name)

        # Add a heading for each portfolio
        doc.add_heading(details.title, level=2)

        # Create a table for the details
        table = doc.add_table(rows=14, cols=2)
        table.style = 'Table Grid'
        table.autofit = False

        # Set the width of the table columns
        for row in table.rows:
            row.cells[0].width = Pt(200)
            row.cells[1].width = Pt(300)

        # Add the details to the table
        details_dict = {
            "Assignment name:": details.title,
            "Approx. value of the contract (in current US$):": details.approximate_contract_value,
            "Country:": details.location,
            "Duration of assignment (months):": details.duration_of_assignment,
            "Name of Client(s):": details.client,
            "Contact Person, Title/Designation, Tel. No./Address:": details.contact,
            "Start Date (month/year):": details.start_date,
            "End Date (month/year):": details.end_date,
            "Total No. of staff-months of the assignment:": details.total_staff_months,
            "No. of professional staff-months provided by your consulting firm/organization or "
            "your sub consultants:": details.total_staff_months,
            "Name of associated Consultants, if any:": "",
            "Name of senior professional staff of your consulting firm/organization involved and "
            "designation and/or functions performed (e.g. Project Director/ Coordinator, "
            "Team Leader):": "",
            "Description of Project:": details.body,
            "Description of actual services provided by your staff within the "
            "assignment:": details.services_listed,
        }

        # Populate the table with the correct details
        for i, (key, value) in enumerate(details_dict.items()):
            cell1 = table.cell(i, 0)
            cell2 = table.cell(i, 1)
            cell1.text = key
            cell2.text = str(value) if value else ""

    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()
